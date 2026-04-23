#!/usr/bin/env python3
"""Generate Vira_Session_Handoff_2026-04-23.docx for the Phase 3 session."""

from docx import Document
from docx.shared import Pt, RGBColor


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x5B, 0x5F, 0x45)  # hemlock
    return h


def add_para(doc, text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x18, 0x1E, 0x14)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


doc = Document()

# Title
title = doc.add_heading("Vira Plants — Session Handoff", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x5B, 0x5F, 0x45)

add_para(doc, "Caretaker mode Phase 3 — invite acceptance (caretaker side)", bold=True)
add_para(doc, "Session: 2026-04-23 · Branch: feature/caretaker-mode-phase-3")

# Summary
add_heading(doc, "What shipped", 1)
bullet(doc, "Migration 004 — accept_garden_invite SECURITY DEFINER RPC (atomic INSERT garden_caretakers + UPDATE garden_invites.accepted_at).")
bullet(doc, "Edge Function accept-invite — JWT decode, email match check, calls the RPC, maps exception messages to error codes, strips owner_email from the response.")
bullet(doc, "caretakerService methods: listPendingInvitesForMe (direct PostgREST + profiles join), acceptInvite (Edge Function), declineInvite (direct PostgREST DELETE).")
bullet(doc, "PendingInviteCard component — Butter Moon card, Vermillion Accept + outlined Decline, inline Vermillion error text, 0.5-opacity + Dismiss-only when expired.")
bullet(doc, "SettingsScreen integration — Pending invitations section at top, section auto-vanishes when empty, pull-to-refresh.")
bullet(doc, "Shared date util — formatFullDate, formatRelative, isInviteExpired extracted to src/utils/formatDate.ts.")
bullet(doc, "questions.md Q1 + Q2 marked resolved.")

# What Sam needs to do
add_heading(doc, "What Sam needs to do before merging / shipping", 1)
add_para(
    doc,
    "1. Apply migration 004 to production.",
    bold=True,
)
add_para(
    doc,
    "Dashboard → SQL Editor. Paste the contents of supabase/migrations/004_accept_invite_rpc.sql. Same pattern as 003 / 003b. Then verify:",
)
add_code(
    doc,
    "SELECT proname, prorettype::regtype\nFROM pg_proc WHERE proname = 'accept_garden_invite';\n-- expect one row",
)
add_code(
    doc,
    "SELECT grantee, privilege_type\nFROM information_schema.routine_privileges\nWHERE routine_name = 'accept_garden_invite';\n-- expect grantee=service_role, privilege=EXECUTE",
)

add_para(
    doc,
    "2. Run the pg_policies sanity check on garden_invites.",
    bold=True,
)
add_code(
    doc,
    "SELECT policyname, cmd, qual\nFROM pg_policies\nWHERE tablename = 'garden_invites'\nORDER BY cmd, policyname;",
)
add_para(
    doc,
    "Expected: 4 rows. invitee_can_view_own_invites (SELECT) and owner_or_invitee_can_delete (DELETE) must both reference auth.jwt() ->> 'email', NOT auth.users. If any policy still references auth.users, STOP and flag.",
)

add_para(
    doc,
    "3. Human interactive tap-through + curl runbook.",
    bold=True,
)
add_para(
    doc,
    "Follow test-results/phase-3-test.md end-to-end. Covers: accept happy path → verify garden_caretakers row, repeat-accept → 409, non-existent invite → 404, email mismatch → 403, decline via DELETE, forced-expired → 410, already_caretaker cleanup. Pre-seeded invite is addressed to sam.morassutti+caretaker1@gmail.com.",
)

add_para(
    doc,
    "4. Review + merge PR.",
    bold=True,
)
add_para(doc, "Branch: feature/caretaker-mode-phase-3. Six commits, ordered:")
bullet(doc, "refactor: extract date/invite helpers to src/utils/formatDate.ts")
bullet(doc, "feat(caretaker): migration 004 accept_garden_invite RPC")
bullet(doc, "feat(caretaker): accept-invite edge function")
bullet(doc, "feat(caretaker): phase 3 service methods and PendingInviteCard")
bullet(doc, "feat(caretaker): phase 3 settings screen integration")
bullet(doc, "docs: phase 3 completion (CLAUDE.md + questions.md + handoff + test-results)")

# Architecture decisions
add_heading(doc, "Architecture decisions worth remembering", 1)
add_para(
    doc,
    "Postgres RPC, not two writes from the Edge Function.",
    bold=True,
)
add_para(
    doc,
    "accept_garden_invite is SECURITY DEFINER, does INSERT garden_caretakers + UPDATE garden_invites.accepted_at in a single transaction. Two-write-from-Edge-Function specifically causes the stale-pending-invite-after-access-granted bug. EXECUTE is granted only to service_role.",
)

add_para(
    doc,
    "Silent accept/decline feedback.",
    bold=True,
)
add_para(
    doc,
    "No toast, no alert. A successful accept just removes the card. The accepted garden becomes visible in Phase 4's home-header garden list — that's the durable confirmation surface. A toast here would be a transient shadow of permanent UI. A TODO(phase-4) in SettingsScreen.handleAcceptInvite marks the exact spot to trigger the future useGardenStore.loadGardens() refresh.",
)

add_para(
    doc,
    "Owner name fallback diverges from elsewhere.",
    bold=True,
)
add_para(
    doc,
    "Caretakers cannot read owner emails via PostgREST, so when profiles.display_name is null the card falls back to 'A Vira gardener' instead of the email-prefix fallback used elsewhere in the app. Accepted as v1 compromise — most users set a display name and the pending window is short (7 days).",
)

add_para(
    doc,
    "Owner/caretaker asymmetry is intentional.",
    bold=True,
)
add_para(
    doc,
    "Owner-side listMyInvites/cancelInvite still goes through the caretaker-invites Edge Function (pre-003b workaround). Caretaker-side listPendingInvitesForMe/declineInvite hit PostgREST directly. Unifying both sides onto direct PostgREST is parked in the Next up list as a post-Phase-6 cleanup. Phase 3 deliberately avoided this refactor to minimize churn.",
)

add_para(
    doc,
    "Owner-seen acceptance badge deferred to Phase 4.",
    bold=True,
)
add_para(
    doc,
    "The owner-facing 'someone accepted your invite' signal (garden_caretakers.owner_seen_at + badge on ManageCaretakersScreen) is parked as a Phase 4 subtask, where it fits the garden-switching work.",
)

# Risks / known caveats
add_heading(doc, "Known caveats", 1)
bullet(
    doc,
    "Migration 004 is NOT yet applied to production. Until Sam runs it via Dashboard SQL Editor, any real accept call returns 500 internal (the RPC is missing). The Edge Function itself is deployed and the JWT/email-mismatch paths work.",
)
bullet(
    doc,
    "The pre-flight pg_policies sanity check was not run by the agent (no DB query access from this environment). Sam should run it in the SQL Editor before the manual tap-through.",
)
bullet(
    doc,
    "The API-level curl tests (Tests 1–7 in phase-3-test.md) are documented as a runbook but were NOT executed end-to-end by the agent — they need a real caretaker JWT, which requires sign-in the agent can't drive.",
)
bullet(
    doc,
    "The agent did run two authenticated-Edge-Function probes (no Authorization header → 401; malformed JWT → 401-could-not-decode-token) which confirm the function's handler runs and --no-verify-jwt is in effect. The platform's ES256 rejection does not fire.",
)

# Verification checks that did pass
add_heading(doc, "Checks that passed in this session", 1)
bullet(doc, "npx tsc --noEmit: zero errors")
bullet(doc, "Hardcoded-hex scan on new files: zero hits")
bullet(doc, "any / @ts-ignore scan on new code: zero hits")
bullet(doc, "Unselectored Zustand store scan on new code: zero hits")
bullet(doc, "iPhone 17 Pro simulator build + launch: succeeded")
bullet(doc, "accept-invite deployed with --no-verify-jwt and confirmed live via agent probes")

# Files touched
add_heading(doc, "Files touched", 1)
add_para(doc, "Added:")
bullet(doc, "supabase/migrations/004_accept_invite_rpc.sql")
bullet(doc, "supabase/functions/accept-invite/index.ts")
bullet(doc, "src/utils/formatDate.ts")
bullet(doc, "src/components/PendingInviteCard.tsx")
bullet(doc, "test-results/phase-3-test.md")
bullet(doc, "scripts/generate_handoff.py")
bullet(doc, "Vira_Session_Handoff_2026-04-23.docx")
add_para(doc, "Modified:")
bullet(doc, "src/services/caretakerService.ts — Phase 3 methods + AcceptedInvite type")
bullet(doc, "src/screens/SettingsScreen.tsx — ScrollView wrap + Pending invitations section")
bullet(doc, "src/screens/ManageCaretakersScreen.tsx — import date helpers from src/utils/formatDate")
bullet(doc, "CLAUDE.md — Phase 3 Done entry, Next up reorder, Implementation Notes + Deployment Learnings")
bullet(doc, "questions.md — Q1 + Q2 moved to Resolved")

doc.save("Vira_Session_Handoff_2026-04-23.docx")
print("Wrote Vira_Session_Handoff_2026-04-23.docx")
